import cv2
import os
from docx import Document
from docx.shared import Inches
from tkinter import Tk, filedialog, Button, Label, StringVar, messagebox

def clear_folder(folder_path):
    """Deletes all files in the specified folder."""
    if os.path.exists(folder_path):
        for file in os.listdir(folder_path):
            file_path = os.path.join(folder_path, file)
            if os.path.isfile(file_path):
                os.remove(file_path)
        print(f"Cleared folder: {folder_path}")
    else:
        os.makedirs(folder_path)
        print(f"Created folder: {folder_path}")

def extract_screenshots(video_path, interval, output_folder):
    """Extracts screenshots from a video at specified intervals."""
    os.makedirs(output_folder, exist_ok=True)
    video = cv2.VideoCapture(video_path)
    fps = int(video.get(cv2.CAP_PROP_FPS))
    frame_interval = interval * fps
    count = 0
    success, frame = video.read()

    while success:
        frame_number = int(video.get(cv2.CAP_PROP_POS_FRAMES))
        if frame_number % frame_interval == 0:
            file_path = os.path.join(output_folder, f'screenshot_{count}.jpg')
            cv2.imwrite(file_path, frame)
            count += 1
        success, frame = video.read()

    video.release()

def create_document_with_screenshots(output_folder, document_name):
    """Creates a Word document with screenshots."""
    doc = Document()
    doc.add_heading('Video Screenshots', level=1)

    for file_name in sorted(os.listdir(output_folder)):
        if file_name.endswith('.jpg'):
            file_path = os.path.join(output_folder, file_name)
            doc.add_paragraph(f'Image: {file_name}')
            doc.add_picture(file_path, width=Inches(6))

    doc.save(document_name)
    print(f"Document saved as {document_name}")

def select_video():
    """Select a video file and process it."""
    video_path = filedialog.askopenfilename(
        title="Select Video File",
        filetypes=[("MP4 Files", "*.mp4"), ("All Files", "*.*")]
    )
    if not video_path:
        messagebox.showwarning("No File", "No video file selected!")
        return

    # Ask for output directory
    output_path = filedialog.askdirectory(
        title="Select Output Folder"
    )
    if not output_path:
        messagebox.showwarning("No Folder", "No output folder selected!")
        return

    clear_folder(output_path)
    extract_screenshots(video_path, interval, output_path)
    document_path = os.path.join(output_path, "VideoScreenshots.docx")
    create_document_with_screenshots(output_path, document_path)

    messagebox.showinfo("Success", f"Processing complete. Document saved at:\n{document_path}")

# Tkinter Frontend
root = Tk()
root.title("Video Screenshot Extractor")
root.geometry("400x200")

interval = 5  # Seconds

Label(root, text="Video Screenshot Extractor", font=("Arial", 16)).pack(pady=10)
Button(root, text="Select Video and Process", command=select_video).pack(pady=10)
Button(root, text="Exit", command=root.quit).pack(pady=10)

root.mainloop()

